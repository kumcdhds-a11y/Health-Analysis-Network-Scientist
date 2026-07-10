# -*- coding: utf-8 -*-
"""역학분석 Word 보고서 생성 (Table1+SMD, crude/adjusted OR, 전체위험, 하위군+P-interaction, forest, E-value)"""
import factory_core as fc
import pandas as pd, numpy as np, json, subprocess, os, io

def _wstat(sub,col):
    s=sub[[col,"wt_pool"]].dropna(); w=s.wt_pool.values; x=s[col].values.astype(float)
    if w.sum()==0: return np.nan,np.nan
    m=np.average(x,weights=w); return m,np.average((x-m)**2,weights=w)

def table1_smd(d,outcome,variables):
    g0=d[d[outcome]==0]; g1=d[d[outcome]==1]; rows=[]
    for v in variables:
        binary=(fc.typ(v)=="b") or (d[v].dropna().nunique()<=2)
        if binary:
            p0,_=_wstat(g0,v); p1,_=_wstat(g1,v); pa,_=_wstat(d,v); pb=(p0+p1)/2
            smd=abs(p1-p0)/np.sqrt(pb*(1-pb)) if 0<pb<1 else np.nan
            rows.append({"Variable":fc.lab(v),"Overall":f"{100*pa:.1f}%","outcome=0":f"{100*p0:.1f}%","outcome=1":f"{100*p1:.1f}%","SMD":f"{smd:.3f}"})
        else:
            m0,v0=_wstat(g0,v); m1,v1=_wstat(g1,v); ma,_=_wstat(d,v); sp=np.sqrt((v0+v1)/2)
            smd=abs(m1-m0)/sp if sp>0 else np.nan
            rows.append({"Variable":fc.lab(v),"Overall":f"{ma:.1f}","outcome=0":f"{m0:.1f}","outcome=1":f"{m1:.1f}","SMD":f"{smd:.3f}"})
    return pd.DataFrame(rows)

def evalue(orr,lo,hi):
    ev=lambda x:round((1/x if x<1 else x)+np.sqrt((1/x if x<1 else x)*((1/x if x<1 else x)-1)),2)
    return ev(orr), ev(hi if orr<1 else lo)

def _forest(adj,labels):
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    a=adj.dropna(subset=["OR"]).copy(); a["name"]=a.term.map(labels)
    fig,ax=plt.subplots(figsize=(6,0.5*len(a)+1))
    y=np.arange(len(a))
    ax.errorbar(a.OR,y,xerr=[a.OR-a.lo,a.hi-a.OR],fmt="o",color="#245",capsize=3)
    ax.axvline(1,color="#999",ls="--"); ax.set_yticks(y); ax.set_yticklabels(a.name,fontsize=8)
    ax.set_xscale("log"); ax.set_xlabel("Adjusted OR (95% CI)"); ax.invert_yaxis(); fig.tight_layout(); return fig

def build_epi_report(dataset, DF, outcome, main_exposure, covariates, subgroups, defs, model, url, use_llm, workdir="."):
    d=fc.apply_definitions(DF,defs); amin=defs["pop"]["age_min"]
    n_total=len(d); d_age=d[d.age>=amin]; n_age=len(d_age)
    # 연속형 → z표준화 항 생성
    def term(v):
        if fc.typ(v)=="c":
            col=f"z_{v}"; d[col]=(d[v]-d[v].mean())/d[v].std(); return col
        return v
    d["age50"]=(d.age>=50).astype(int)  # 하위군용 연령 이분
    ex_t=term(main_exposure); cov_t=[term(c) for c in covariates]
    labels={ex_t:fc.lab(main_exposure),**{term(c):fc.lab(c) for c in covariates}}
    # 인과분석용 이분 노출(연속이면 가중 중앙값 기준 이분), RCS용 연속항
    def wmedian(x,w):
        s=pd.DataFrame({"x":x,"w":w}).dropna().sort_values("x"); c=s.w.cumsum()
        return s.x[c>=s.w.sum()/2].iloc[0]
    if fc.typ(main_exposure)=="c":
        med=wmedian(d[main_exposure],d["wt_pool"]); d["xbin"]=(d[main_exposure]>med).astype("Int64")
        exposure_bin="xbin"; exposure_cont=ex_t; xbin_label=f"{fc.lab(main_exposure)} high (above median)"
    else:
        exposure_bin=main_exposure; exposure_cont=""; xbin_label=fc.lab(main_exposure)
    need=list(dict.fromkeys([outcome,ex_t,exposure_bin]+cov_t+subgroups+["age","kstrata","psu","wt_pool"]))
    ana=d[(d.age>=amin)&(d.wt_pool>0)&d.wt_pool.notna()&d.kstrata.notna()&d.psu.notna()].dropna(subset=[outcome]).copy()
    n_ana=len(ana)
    ana[[c for c in need if c in ana.columns]].to_csv(os.path.join(workdir,"epi_analytic.csv"),index=False)
    json.dump({"outcome":outcome,"exposure":ex_t,"cov":cov_t,"subgroups":subgroups,
               "exposure_bin":exposure_bin,"exposure_cont":exposure_cont},
              open(os.path.join(workdir,"epi_config.json"),"w"),ensure_ascii=False)
    r=subprocess.run(["Rscript","epi.R"],cwd=workdir,capture_output=True,text=True)
    if r.returncode!=0: raise RuntimeError(r.stderr)
    ra=subprocess.run(["Rscript","epi_adv.R"],cwd=workdir,capture_output=True,text=True)
    if ra.returncode!=0: raise RuntimeError("epi_adv: "+ra.stderr)
    vif=pd.read_csv(os.path.join(workdir,"epi_vif.csv")); rcs=pd.read_csv(os.path.join(workdir,"epi_rcs.csv"))
    love=pd.read_csv(os.path.join(workdir,"epi_love.csv")); tab6=pd.read_csv(os.path.join(workdir,"epi_table6.csv"))
    crude=pd.read_csv(os.path.join(workdir,"epi_crude.csv")); adj=pd.read_csv(os.path.join(workdir,"epi_adj.csv"))
    risk=pd.read_csv(os.path.join(workdir,"epi_risk.csv")); sub=pd.read_csv(os.path.join(workdir,"epi_sub.csv"))
    t1=table1_smd(ana,outcome,[main_exposure]+covariates)
    me=adj[adj.term==ex_t].iloc[0]; evp,evc=evalue(me.OR,me.lo,me.hi)
    fig=_forest(adj,labels)
    # LLM 해석 (선택)
    facts=(f"Outcome {fc.lab(outcome)}. Main exposure {fc.lab(main_exposure)} adjusted OR {me.OR:.2f} "
           f"(95% CI {me.lo:.2f}-{me.hi:.2f}). E-value {evp} (CI {evc}). Weighted risk {risk.weighted_risk_pct[0]:.1f}%.")
    interp=""
    if use_llm:
        try: interp=fc.ollama_chat(f"{fc.STYLE}\nWrite the Results paragraph of a research paper in English from the epidemiologic results below. Do not change any numbers.\n{facts}",model,url).strip()
        except Exception: interp=""
    if not interp:
        interp=(f"In survey-weighted logistic regression, {fc.lab(main_exposure)} showed an adjusted odds ratio of "
                f"{me.OR:.2f} (95% CI {me.lo:.2f} to {me.hi:.2f}) for {fc.lab(outcome)}. The E-value was {evp}, "
                f"indicating the minimum strength of unmeasured confounding on the odds-ratio scale that could explain the estimate.")
    # ── docx ──
    from docx import Document
    from docx.shared import Inches
    doc=Document(); doc.add_heading("Epidemiologic Analysis Report",0)
    doc.add_paragraph(f"Outcome: {fc.lab(outcome)} · {dataset} · survey-weighted logistic regression")
    doc.add_paragraph(f"[LOGISTIC] outcome = {outcome}")
    doc.add_heading("Figure 1. Study population selection flow",1)
    doc.add_paragraph(f"Total {n_total:,} → age ≥{amin} years {n_age:,} → complete design and outcome data {n_ana:,} (analytic sample).")
    def add_tab(title, df, note=""):
        doc.add_heading(title,1)
        if note: doc.add_paragraph(note)
        cols=list(df.columns); tb=doc.add_table(rows=1,cols=len(cols)); tb.style="Light Grid Accent 1"
        for j,c in enumerate(cols): tb.rows[0].cells[j].text=str(c)
        for _,row in df.iterrows():
            cc=tb.add_row().cells
            for j,c in enumerate(cols): cc[j].text=str(row[c])
    add_tab("Table 1. Baseline characteristics (weighted, with SMD)", t1)
    cr=crude.copy(); cr["OR (95% CI)"]=cr.apply(lambda x:f"{x.OR:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.OR) else "-",axis=1)
    cr["Variable"]=cr.term.map(labels); cr["p"]=cr.p.map(lambda p:"<0.001" if pd.notna(p) and p<0.001 else (f"{p:.3f}" if pd.notna(p) else ""))
    add_tab("Table 2. Crude OR (univariable)", cr[["Variable","OR (95% CI)","p"]])
    aj=adj.copy(); aj["OR (95% CI)"]=aj.apply(lambda x:f"{x.OR:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.OR) else "-",axis=1)
    aj["Variable"]=aj.term.map(labels); aj["p"]=aj.p.map(lambda p:"<0.001" if pd.notna(p) and p<0.001 else (f"{p:.3f}" if pd.notna(p) else ""))
    add_tab("Table 3. Adjusted OR (multivariable)", aj[["Variable","OR (95% CI)","p"]])
    add_tab("Table 4. Overall risk (weighted)", pd.DataFrame({"N":[n_ana],"Weighted risk (%)":[f"{risk.weighted_risk_pct[0]:.1f}"]}))
    SUBLAB={"men":"Sex","age50":"Age group"}
    def lvlab(s,lv):
        if s=="men": return "Men" if lv==1 else "Women"
        if s=="age50": return "≥50" if lv==1 else "<50"
        return str(lv)
    sg=sub.copy(); sg["OR (95% CI)"]=sg.apply(lambda x:f"{x.OR:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.OR) else "-",axis=1)
    sg["Subgroup"]=sg.subgroup.map(lambda s:SUBLAB.get(s, fc.lab(s) if s in fc.VARS else s))
    sg["Level"]=sg.apply(lambda r:lvlab(r.subgroup,r.level),axis=1)
    sg["P-interaction"]=sg.p_int.map(lambda p:"" if pd.isna(p) else ("<0.001" if p<0.001 else f"{p:.3f}"))
    add_tab(f"Table 5. Subgroup analysis — {fc.lab(main_exposure)} effect with P for interaction",
            sg[["Subgroup","Level","OR (95% CI)","P-interaction"]])
    doc.add_heading("Adjusted OR forest",1)
    bio=io.BytesIO(); fig.savefig(bio,format="png",dpi=110,bbox_inches="tight"); bio.seek(0)
    import matplotlib.pyplot as plt; plt.close(fig)
    doc.add_picture(bio,width=Inches(5.5))
    doc.add_heading("E-value",1)
    doc.add_paragraph(f"E-value for the main exposure {fc.lab(main_exposure)} (adjusted OR {me.OR:.2f}) = {evp} (confidence-interval limit {evc}).")
    doc.add_heading("Interpretation",1); doc.add_paragraph(interp)
    # ── Table 6: 노출효과 종합 (Crude·Min-adj·Full-adj·IPTW·G-comp·AIPW) ──
    doc.add_heading(f"Table 6. Main exposure effect summary — {xbin_label}",1)
    t6=tab6.copy(); t6["Estimate (95% CI)"]=t6.apply(
        lambda x:f"{x.estimate:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.lo) else f"{x.estimate:.3f}",axis=1)
    add_tab("", t6[["method","scale","Estimate (95% CI)"]].rename(columns={"method":"Method","scale":"Scale"}))
    # ── Love plot (IPTW 전후 SMD) ──
    doc.add_heading("Love plot — covariate balance (|SMD| before/after IPTW)",1)
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    lv=love.dropna(); figl,axl=plt.subplots(figsize=(6,0.4*len(lv)+1)); yy=np.arange(len(lv))
    axl.scatter(lv.before,yy,label="before",color="#A32D2D"); axl.scatter(lv.after,yy,label="after (IPTW)",color="#0F6E56")
    axl.axvline(0.1,color="#999",ls="--"); axl.set_yticks(yy); axl.set_yticklabels([labels.get(c,c) for c in lv.covariate],fontsize=8)
    axl.set_xlabel("|SMD|"); axl.legend(fontsize=8); axl.invert_yaxis(); figl.tight_layout()
    bl=io.BytesIO(); figl.savefig(bl,format="png",dpi=110,bbox_inches="tight"); bl.seek(0); plt.close(figl)
    doc.add_picture(bl,width=Inches(5.5))
    # ── VIF ──
    doc.add_heading("VIF (multicollinearity)",1)
    vt=vif.copy(); vt["Variable"]=vt.term.map(lambda t:labels.get(t,t)); vt["VIF"]=vt.VIF
    add_tab("", vt[["Variable","VIF"]])
    # ── RCS 비선형성 ──
    pnl=rcs.value.iloc[0]
    doc.add_heading("Restricted cubic spline nonlinearity test",1)
    doc.add_paragraph("Not applicable (binary exposure)." if pd.isna(pnl) else
        f"RCS (3 knots) nonlinearity P = {'<0.001' if pnl<0.001 else f'{pnl:.3f}'} "
        f"({'nonlinearity significant' if pnl<0.05 else 'linearity not rejected'}).")
    buf=io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── 생존분석 보고서 (설계가중 Cox svycoxph + KM + 발생률) ──
def build_survival_report(dataset, DF, time_col, event_col, main_exposure, covariates, defs, label,
                          model, url, use_llm, before_fu_years=None, workdir="."):
    """label 예: 'survival death','survival MACE'. before_fu_years 지정 시 그 이전 사건 제외(사전 추적 사망 처리)."""
    d=fc.apply_definitions(DF,defs); amin=defs["pop"]["age_min"]
    if time_col not in d.columns or event_col not in d.columns:
        raise RuntimeError(f"Follow-up/event variables ({time_col},{event_col}) not found. Please specify cohort variables.")
    def term(v):
        if fc.typ(v)=="c": d[f"z_{v}"]=(d[v]-d[v].mean())/d[v].std(); return f"z_{v}"
        return v
    ex_t=term(main_exposure); cov_t=[term(c) for c in covariates]
    labels={ex_t:fc.lab(main_exposure),**{term(c):fc.lab(c) for c in covariates}}
    def wmedian(x,w):
        s=pd.DataFrame({"x":x,"w":w}).dropna().sort_values("x"); c=s.w.cumsum(); return s.x[c>=s.w.sum()/2].iloc[0]
    if fc.typ(main_exposure)=="c":
        d["xbin"]=(d[main_exposure]>wmedian(d[main_exposure],d["wt_pool"])).astype("Int64"); exposure_bin="xbin"; xbl=f"{fc.lab(main_exposure)} high"
    else: exposure_bin=main_exposure; xbl=fc.lab(main_exposure)
    d["ftime"]=pd.to_numeric(d[time_col],errors="coerce"); d["fevent"]=pd.to_numeric(d[event_col],errors="coerce")
    ana=d[(d.age>=amin)&(d.wt_pool>0)&d.wt_pool.notna()&d.kstrata.notna()&d.psu.notna()&d.ftime.notna()&d.fevent.notna()&(d.ftime>0)].copy()
    note=""
    if before_fu_years is not None:  # 사전 추적 사망: 기준시점 이전(짧은 추적) 사건 제외 sensitivity
        n0=len(ana); ana=ana[~((ana.fevent==1)&(ana.ftime<before_fu_years))]; note=f"Excluded {n0-len(ana)} events occurring before {before_fu_years} years of follow-up (pre-follow-up handling)."
    need=list(dict.fromkeys(["ftime","fevent",ex_t,exposure_bin]+cov_t+["age","kstrata","psu","wt_pool"]))
    ana[[c for c in need if c in ana.columns]].to_csv(f"{workdir}/surv_analytic.csv",index=False)
    import json,subprocess,io
    json.dump({"time":"ftime","event":"fevent","exposure":ex_t,"exposure_bin":exposure_bin,"cov":cov_t},open(f"{workdir}/surv_config.json","w"),ensure_ascii=False)
    r=subprocess.run(["Rscript","epi_surv.R"],cwd=workdir,capture_output=True,text=True)
    if r.returncode!=0: raise RuntimeError("epi_surv: "+r.stderr)
    crude=pd.read_csv(f"{workdir}/surv_crude.csv"); adj=pd.read_csv(f"{workdir}/surv_adj.csv")
    km=pd.read_csv(f"{workdir}/surv_km.csv"); inc=pd.read_csv(f"{workdir}/surv_inc.csv")
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    figk,axk=plt.subplots(figsize=(6,4))
    for g,sub in km.groupby("group"):
        axk.step(sub.time,sub.surv,where="post",label=f"{main_exposure}={g}")
    axk.set_xlabel("Time"); axk.set_ylabel("Survival probability"); axk.set_ylim(0,1); axk.legend(fontsize=8); axk.set_title(f"KM — {label}"); figk.tight_layout()
    me=adj[adj.term==ex_t]; hrtxt = f"{me.HR.iloc[0]:.2f} ({me.lo.iloc[0]:.2f}-{me.hi.iloc[0]:.2f})" if len(me) else "-"
    facts=f"{label}: adjusted hazard ratio for {fc.lab(main_exposure)} {hrtxt}. Incidence {inc.rate_per_1000py.iloc[0]}/1000 person-years."
    interp=""
    if use_llm:
        try: interp=fc.ollama_chat(f"{fc.STYLE}\nWrite the survival-analysis results as English paper sentences. Do not change any numbers.\n{facts}",model,url).strip()
        except Exception: interp=""
    if not interp: interp=f"In survey-weighted Cox regression, {fc.lab(main_exposure)} was associated with {label} with an adjusted hazard ratio of {hrtxt}."
    from docx import Document; from docx.shared import Inches
    doc=Document(); doc.add_heading(f"Survival Analysis Report — {label}",0)
    doc.add_paragraph(f"{dataset} · survey-weighted Cox regression (svycoxph) · outcome = {label}")
    if note: doc.add_paragraph(note)
    def tab(title,df):
        doc.add_heading(title,1); cols=list(df.columns); tb=doc.add_table(rows=1,cols=len(cols)); tb.style="Light Grid Accent 1"
        for j,c in enumerate(cols): tb.rows[0].cells[j].text=str(c)
        for _,rw in df.iterrows():
            cc=tb.add_row().cells
            for j,c in enumerate(cols): cc[j].text=str(rw[c])
    doc.add_heading("Incidence (survey-weighted)",1); doc.add_paragraph(f"Events {int(inc.events.iloc[0]):,}, person-time {int(inc.person_time.iloc[0]):,}, incidence rate {inc.rate_per_1000py.iloc[0]} per 1000 person-years")
    doc.add_heading("Kaplan-Meier survival curves",1); b=io.BytesIO(); figk.savefig(b,format="png",dpi=110,bbox_inches="tight"); b.seek(0); plt.close(figk); doc.add_picture(b,width=Inches(5.5))
    cr=crude.copy(); cr["Variable"]=cr.term.map(lambda t:labels.get(t,t)); cr["HR (95% CI)"]=cr.apply(lambda x:f"{x.HR:.2f} ({x.lo:.2f}-{x.hi:.2f})",axis=1)
    tab("Crude HR (Cox)", cr[["Variable","HR (95% CI)","p"]])
    aj=adj.copy(); aj["Variable"]=aj.term.map(lambda t:labels.get(t,t)); aj["HR (95% CI)"]=aj.apply(lambda x:f"{x.HR:.2f} ({x.lo:.2f}-{x.hi:.2f})",axis=1)
    tab("Adjusted HR (Cox)", aj[["Variable","HR (95% CI)","p"]])
    doc.add_heading("Interpretation",1); doc.add_paragraph(interp)
    buf=io.BytesIO(); doc.save(buf); return buf.getvalue()
